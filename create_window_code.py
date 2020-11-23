from gui.CreateWindowGUI import CreateWindowGUI
from task_widget_code import TaskWidget
from docx import Document
from task_cod import TaskGenerator
from PyQt5.QtWidgets import QFileDialog, QWidget, QVBoxLayout
import subprocess
import os
import sys

sys.excepthook = lambda *a: sys.__excepthook__(*a)


class CreateWindow(QWidget, CreateWindowGUI):
    def __init__(self):
        self.task_widgets = []
        super().__init__()
        self.setupUi(self)
        self.add_button.clicked.connect(self.add)
        self.generate_button.clicked.connect(self.generate)
        self.generate_button.setEnabled(False)
        self.tasks_layout = QVBoxLayout()
        self.qw = QWidget()
        self.qw.setFixedHeight(150)
        self.n = 0

        self.qw.setLayout(self.tasks_layout)
        self.task_area.setWidget(self.qw)

    def open_file(self):
        self.task_widgets.clear()
        for widget in self.task_widgets:
            widget.hide()
        directory = QFileDialog.getExistingDirectory()
        print(directory)
        setting_file = open(directory + '/setting.txt', 'r', encoding='utf-8')
        settings = setting_file.readlines()
        self.name_edit.setText(settings.pop(0).strip())
        self.variants_count_edit.setValue(int(settings.pop(0).strip()))
        for task in settings:
            self.add()
            self.task_widgets[-1].set_parametrs(*task.split(';'))
        setting_file.close()

    def add(self):
        self.n += 1
        task = TaskWidget(self)
        self.task_widgets.append(task)
        self.tasks_layout.addWidget(task)
        self.qw.setFixedHeight(10 + 140 * self.n)
        self.generate_button.setEnabled(True)

    def generate(self):

        def save():
            setting_file = open(directory + '/setting.txt', 'w', encoding='utf-8')
            settings = []
            settings.append(self.name_edit.text())
            settings.append(self.variants_count_edit.text())
            for task in self.task_widgets:
                settings.append(str(task.get_id()) + ';' + task.count_edit.text())
            subprocess.call(['attrib', '+h', directory + '/setting.txt'])
            setting_file.write('\n'.join(settings))
            setting_file.close()

        def generate_document():
            tasks = Document()
            tasks.add_heading(self.name_edit.text() + f' Вариант-{variant}', 0)
            number = 1
            for task_widget in self.task_widgets:
                for _ in range(task_widget.get_count()):
                    task_gen = TaskGenerator(task_widget.get_pattern())
                    p = tasks.add_paragraph(f'№{number} ' + task_gen.get_text()[0])
                    p.alignment = 3
                    par.add_run(f'\n№{number} - ' + str(task_gen.get_text()[1]))
                    number += 1
            tasks.save(directory + '/' + self.name_edit.text() + f'_вариант-{variant}.docx')

        file_dialog = QFileDialog()

        directory = \
            file_dialog.getSaveFileName(self, self.name_edit.text(), 'имя работы',
                                        'Все файлы (*)')[0]
        os.mkdir(directory)

        answers = Document()
        for variant in range(1, int(self.variants_count_edit.text()) + 1):
            answers.add_heading(f'Вариант-{variant}', 0)
            par = answers.add_paragraph()
            generate_document()
        answers.save(directory + '/' + self.name_edit.text() + '_ответы.docx')
        save()
        self.close()
